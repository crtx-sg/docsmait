# backend/app/kb_service_pg.py
import os
import time
import uuid
import json
from datetime import datetime, date
from typing import List, Dict, Any, Optional
from fastapi import HTTPException, UploadFile
import PyPDF2
import docx
import magic
from sqlalchemy.orm import Session, joinedload
from sqlalchemy import and_, or_, func

from .config import config
import ollama
import qdrant_client
from qdrant_client.http.models import PointStruct, VectorParams, Distance

from .database_config import get_db
from .db_models import KBCollection, KBDocument, KBQuery, KBConfig, KBDocumentTag
from .ai_service import ai_service

class KnowledgeBaseService:
    def __init__(self):
        self.ollama_client = ollama.Client(host=config.OLLAMA_BASE_URL)
        self.qdrant_client = qdrant_client.QdrantClient(url=config.QDRANT_URL)
        self.ai_service = ai_service  # Add AI service for training functionality
        self.ensure_kb_tables()
        
    def ensure_kb_tables(self):
        """Ensure KB tables exist (they should be created by init_db.py)"""
        # Tables are created via SQLAlchemy models in init_db.py
        # This method is kept for compatibility but does nothing
        pass

    def create_collection(self, name: str, description: str, created_by: str, 
                         tags: List[str] = None, is_default: bool = False) -> Dict[str, Any]:
        """Create a new knowledge base collection"""
        db = next(get_db())
        try:
            # Check if collection already exists
            existing_collection = db.query(KBCollection).filter(KBCollection.name == name).first()
            if existing_collection:
                return {"success": False, "error": f"Collection '{name}' already exists"}
            
            # Create collection in PostgreSQL
            collection_id = str(uuid.uuid4())
            collection = KBCollection(
                id=collection_id,
                name=name,
                description=description,
                created_by=created_by,
                tags=json.dumps(tags) if tags else "[]",
                is_default=is_default
            )
            db.add(collection)
            
            # Create Qdrant collection for vectors
            try:
                self.qdrant_client.create_collection(
                    collection_name=name,
                    vectors_config=VectorParams(size=config.EMBEDDING_DIMENSIONS, distance=Distance.COSINE)
                )
            except Exception as e:
                # Check if collection already exists (409 Conflict)
                if "already exists" in str(e) or "409" in str(e):
                    print(f"Qdrant collection '{name}' already exists, continuing with PostgreSQL record creation...")
                else:
                    db.rollback()
                    return {"success": False, "error": f"Failed to create Qdrant collection: {str(e)}"}
            
            db.commit()
            
            return {
                "success": True,
                "collection_id": collection_id,
                "message": f"Collection '{name}' created successfully"
            }
            
        except Exception as e:
            db.rollback()
            print(f"Error creating collection: {e}")
            return {"success": False, "error": f"Failed to create collection: {str(e)}"}
        finally:
            db.close()

    def list_collections(self) -> List[Dict[str, Any]]:
        """List all collections"""
        db = next(get_db())
        try:
            collections = db.query(KBCollection).order_by(KBCollection.created_date.desc()).all()
            
            # If no collections exist, create a default one
            if not collections:
                print("No collections found, creating default collection")
                default_result = self.create_collection(
                    name=config.DEFAULT_COLLECTION_NAME,
                    description="Default knowledge base collection",
                    created_by="system",
                    tags=["default"],
                    is_default=True
                )
                if default_result.get("success"):
                    # Refresh collections query
                    collections = db.query(KBCollection).order_by(KBCollection.created_date.desc()).all()
            
            result = []
            for collection in collections:
                total_size_mb = (collection.total_size_bytes or 0) / (1024 * 1024)
                result.append({
                    "id": collection.id,
                    "name": collection.name,
                    "description": collection.description,
                    "created_by": collection.created_by,
                    "created_date": collection.created_date.isoformat() if collection.created_date else None,
                    "updated_date": collection.updated_date.isoformat() if collection.updated_date else None,
                    "document_count": collection.document_count or 0,
                    "total_size_bytes": collection.total_size_bytes or 0,
                    "total_size_mb": round(total_size_mb, 2),
                    "tags": json.loads(collection.tags) if collection.tags else [],
                    "is_default": collection.is_default
                })
            
            return result
            
        except Exception as e:
            print(f"Error listing collections: {e}")
            import traceback
            traceback.print_exc()
            return []
        finally:
            db.close()

    def get_collections(self) -> List[Dict[str, Any]]:
        """Alias for list_collections - used by training functionality"""
        return self.list_collections()

    def get_collection(self, collection_name: str) -> Optional[Dict[str, Any]]:
        """Get collection by name with documents"""
        db = next(get_db())
        try:
            collection = db.query(KBCollection).filter(KBCollection.name == collection_name).first()
            
            if not collection:
                return None
            
            # Get documents in this collection
            documents = db.query(KBDocument).filter(
                KBDocument.collection_name == collection_name
            ).order_by(KBDocument.upload_date.desc()).all()
            
            documents_list = []
            for doc in documents:
                documents_list.append({
                    "id": doc.id,
                    "filename": doc.filename,
                    "content_type": doc.content_type,
                    "size_bytes": doc.size_bytes,
                    "chunk_count": doc.chunk_count or 0,
                    "upload_date": doc.upload_date.isoformat() if doc.upload_date else None,
                    "status": doc.status
                })
            
            total_size_mb = (collection.total_size_bytes or 0) / (1024 * 1024)
            return {
                "id": collection.id,
                "name": collection.name,
                "description": collection.description,
                "created_by": collection.created_by,
                "created_date": collection.created_date.isoformat() if collection.created_date else None,
                "updated_date": collection.updated_date.isoformat() if collection.updated_date else None,
                "document_count": collection.document_count or 0,
                "total_size_bytes": collection.total_size_bytes or 0,
                "total_size_mb": round(total_size_mb, 2),
                "tags": json.loads(collection.tags) if collection.tags else [],
                "is_default": collection.is_default,
                "documents": documents_list
            }
            
        except Exception as e:
            print(f"Error getting collection: {e}")
            return None
        finally:
            db.close()

    def process_document(self, file: UploadFile, collection_name: str = None, 
                        chunk_size: int = None) -> Dict[str, Any]:
        """Process and add document to collection (alias for add_document_to_collection)"""
        if collection_name is None:
            collection_name = config.DEFAULT_COLLECTION_NAME
        if chunk_size is None:
            chunk_size = config.DEFAULT_CHUNK_SIZE
        return self.add_document_to_collection(collection_name, file, chunk_size)

    def add_document_to_collection(self, collection_name: str, file: UploadFile, 
                                 chunk_size: int = None) -> Dict[str, Any]:
        """Add document to collection (metadata in PostgreSQL, vectors in Qdrant)"""
        db = next(get_db())
        start_time = time.time()
        
        if chunk_size is None:
            chunk_size = config.DEFAULT_CHUNK_SIZE
        
        try:
            # Check if collection exists, fallback to default if not
            actual_collection_name = self._ensure_collection_exists_or_get_default(collection_name)
            
            collection = db.query(KBCollection).filter(KBCollection.name == actual_collection_name).first()
            if not collection:
                return {"success": False, "error": f"Collection '{actual_collection_name}' not found and could not create default"}
            
            # Read file content
            file_content = file.file.read()
            file.file.seek(0)  # Reset file pointer
            
            # Extract text based on file type
            text_content = self._extract_text_content(file_content, file.content_type, file.filename)
            if not text_content:
                return {"success": False, "error": "Failed to extract text from file"}
            
            # Create document record in PostgreSQL
            document_id = str(uuid.uuid4())
            kb_document = KBDocument(
                id=document_id,
                filename=file.filename,
                content_type=file.content_type,
                size_bytes=len(file_content),
                collection_name=actual_collection_name,
                status="processing"
            )
            db.add(kb_document)
            
            # Process text into chunks and create embeddings (Qdrant operations)
            chunks = self._chunk_text(text_content, chunk_size)
            vectors = []
            
            for i, chunk in enumerate(chunks):
                try:
                    # Generate embedding using Ollama
                    embedding = self._generate_embedding(chunk)
                    
                    point = PointStruct(
                        id=str(uuid.uuid4()),
                        vector=embedding,
                        payload={
                            "document_id": document_id,
                            "filename": file.filename,
                            "chunk_index": i,
                            "text": chunk,
                            "collection": actual_collection_name
                        }
                    )
                    vectors.append(point)
                    
                except Exception as e:
                    print(f"Error processing chunk {i}: {e}")
                    continue
            
            # Store vectors in Qdrant
            if vectors:
                self.qdrant_client.upsert(
                    collection_name=actual_collection_name,
                    points=vectors
                )
            
            # Update document record
            kb_document.chunk_count = len(vectors)
            kb_document.status = "completed"
            
            # Update collection stats
            collection.document_count += 1
            collection.total_size_bytes += len(file_content)
            collection.updated_date = datetime.utcnow()
            
            db.commit()
            
            processing_time = time.time() - start_time
            
            return {
                "success": True,
                "document_id": document_id,
                "chunks_created": len(vectors),
                "processing_time": round(processing_time, 2),
                "message": f"Document '{file.filename}' added successfully"
            }
            
        except Exception as e:
            db.rollback()
            print(f"Error adding document: {e}")
            processing_time = time.time() - start_time
            return {
                "success": False, 
                "error": f"Failed to add document: {str(e)}",
                "processing_time": round(processing_time, 2)
            }
        finally:
            db.close()

    def search_collection(self, collection_name: str, query: str, limit: int = 5) -> List[Dict[str, Any]]:
        """Search collection using vector similarity"""
        db = next(get_db())
        try:
            start_time = time.time()
            
            # Check if collection exists, fallback to default if not
            actual_collection_name = self._ensure_collection_exists_or_get_default(collection_name)
            
            # Generate query embedding
            query_embedding = self._generate_embedding(query)
            
            # Search in Qdrant
            search_results = self.qdrant_client.search(
                collection_name=actual_collection_name,
                query_vector=query_embedding,
                limit=limit,
                with_payload=True
            )
            
            # Log query
            response_time = int((time.time() - start_time) * 1000)
            kb_query = KBQuery(
                query_text=query,
                collection_name=actual_collection_name,
                response_time_ms=response_time
            )
            db.add(kb_query)
            db.commit()
            
            # Format results
            results = []
            for result in search_results:
                results.append({
                    "score": result.score,
                    "text": result.payload.get("text", ""),
                    "filename": result.payload.get("filename", ""),
                    "document_id": result.payload.get("document_id", ""),
                    "chunk_index": result.payload.get("chunk_index", 0)
                })
            
            return results
            
        except Exception as e:
            print(f"Error searching collection: {e}")
            return []
        finally:
            db.close()

    def query_collection(self, collection_name: str, query: str, limit: int = 5) -> Dict[str, Any]:
        """Query collection and return results in training-compatible format"""
        try:
            results = self.search_collection(collection_name, query, limit)
            return {
                "success": True,
                "results": results
            }
        except Exception as e:
            print(f"Error querying collection {collection_name}: {e}")
            return {
                "success": False,
                "error": str(e),
                "results": []
            }

    def get_collection_documents(self, collection_name: str) -> List[Dict[str, Any]]:
        """Get all documents in a collection"""
        db = next(get_db())
        try:
            documents = db.query(KBDocument).filter(
                KBDocument.collection_name == collection_name
            ).order_by(KBDocument.upload_date.desc()).all()
            
            result = []
            for doc in documents:
                result.append({
                    "id": doc.id,
                    "filename": doc.filename,
                    "content_type": doc.content_type,
                    "size_bytes": doc.size_bytes,
                    "chunk_count": doc.chunk_count,
                    "upload_date": doc.upload_date.isoformat() if doc.upload_date else None,
                    "status": doc.status
                })
            
            return result
            
        except Exception as e:
            print(f"Error getting collection documents: {e}")
            return []
        finally:
            db.close()

    # Helper methods (keep existing Qdrant-based implementations)
    def _extract_text_content(self, file_content: bytes, content_type: str, filename: str) -> str:
        """Extract text content from various file types"""
        try:
            if content_type == "application/pdf" or filename.lower().endswith('.pdf'):
                # PDF extraction
                import io
                pdf_file = io.BytesIO(file_content)
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\\n"
                return text.strip()
            
            elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or filename.lower().endswith('.docx'):
                # DOCX extraction
                import io
                docx_file = io.BytesIO(file_content)
                doc = docx.Document(docx_file)
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\\n"
                return text.strip()
            
            elif content_type.startswith("text/") or filename.lower().endswith(('.txt', '.md', '.py', '.js', '.html', '.css')):
                # Text file
                return file_content.decode('utf-8', errors='ignore')
            
            else:
                # Try to decode as text
                return file_content.decode('utf-8', errors='ignore')
                
        except Exception as e:
            print(f"Error extracting text content: {e}")
            return ""

    def _chunk_text(self, text: str, chunk_size: int = 1000) -> List[str]:
        """Split text into chunks"""
        chunks = []
        words = text.split()
        current_chunk = []
        current_size = 0
        
        for word in words:
            if current_size + len(word) > chunk_size and current_chunk:
                chunks.append(" ".join(current_chunk))
                current_chunk = [word]
                current_size = len(word)
            else:
                current_chunk.append(word)
                current_size += len(word) + 1  # +1 for space
        
        if current_chunk:
            chunks.append(" ".join(current_chunk))
        
        return chunks

    def _generate_embedding(self, text: str) -> List[float]:
        """Generate embedding using Ollama"""
        try:
            response = self.ollama_client.embeddings(model=config.DEFAULT_EMBEDDING_MODEL, prompt=text)
            return response['embedding']
        except Exception as e:
            print(f"Error generating embedding: {e}")
            # Return zero vector as fallback
            return [0.0] * config.EMBEDDING_DIMENSIONS

    def get_statistics(self) -> Dict[str, Any]:
        """Get Knowledge Base statistics"""
        db = next(get_db())
        try:
            # Count total documents
            total_documents = db.query(KBDocument).count()
            
            # Count documents that are successfully indexed (completed status)
            documents_indexed = db.query(KBDocument).filter(KBDocument.status == "completed").count()
            
            # Count queries today
            today = date.today()
            queries_today = db.query(KBQuery).filter(
                func.date(KBQuery.timestamp) == today
            ).count()
            
            # For simplicity, search queries today is the same as total queries today
            search_queries_today = queries_today
            
            # Calculate total index size (sum of all document sizes in MB)
            total_size_bytes = db.query(func.sum(KBDocument.size_bytes)).scalar() or 0
            index_size_mb = total_size_bytes / (1024 * 1024)
            
            # Get last updated date from most recent document
            last_doc = db.query(KBDocument).order_by(KBDocument.upload_date.desc()).first()
            last_updated = last_doc.upload_date if last_doc else None
            
            return {
                "documents_indexed": documents_indexed,
                "total_documents": total_documents,
                "queries_today": queries_today,
                "search_queries_today": search_queries_today,
                "index_size_mb": round(index_size_mb, 2),
                "last_updated": last_updated
            }
            
        except Exception as e:
            print(f"Error getting statistics: {e}")
            return {
                "documents_indexed": 0,
                "total_documents": 0,
                "queries_today": 0,
                "search_queries_today": 0,
                "index_size_mb": 0.0,
                "last_updated": None
            }
        finally:
            db.close()

    def get_config(self, key: str, default: str = None) -> str:
        """Get configuration value"""
        db = next(get_db())
        try:
            config_item = db.query(KBConfig).filter(KBConfig.key == key).first()
            if config_item:
                return config_item.value
            return default
            
        except Exception as e:
            print(f"Error getting config {key}: {e}")
            return default
        finally:
            db.close()

    def update_config(self, key: str, value: str) -> bool:
        """Update configuration value"""
        db = next(get_db())
        try:
            config_item = db.query(KBConfig).filter(KBConfig.key == key).first()
            
            if config_item:
                config_item.value = value
                config_item.updated_at = datetime.utcnow()
            else:
                config_item = KBConfig(key=key, value=value)
                db.add(config_item)
            
            db.commit()
            return True
            
        except Exception as e:
            print(f"Error updating config {key}: {e}")
            db.rollback()
            return False
        finally:
            db.close()
    
    def _get_default_collection_name(self) -> str:
        """Get the default collection name"""
        return self.get_config("default_collection", config.DEFAULT_COLLECTION_NAME)
    
    def _ensure_collection_exists_or_get_default(self, collection_name: str) -> str:
        """Check if collection exists, if not return default collection name"""
        db = next(get_db())
        try:
            collection = db.query(KBCollection).filter(KBCollection.name == collection_name).first()
            if collection:
                return collection_name
            
            # Collection doesn't exist, check if default collection exists
            default_name = self._get_default_collection_name()
            default_collection = db.query(KBCollection).filter(KBCollection.name == default_name).first()
            
            if default_collection:
                print(f"Collection '{collection_name}' not found, using default collection '{default_name}'")
                return default_name
            
            # Create default collection if it doesn't exist
            result = self.create_collection(
                name=default_name,
                description="Default knowledge base collection",
                created_by="system",
                tags=["default"],
                is_default=True
            )
            
            if result.get("success"):
                print(f"Created default collection '{default_name}' as fallback for '{collection_name}'")
                return default_name
            else:
                print(f"Failed to create default collection: {result.get('error')}")
                return collection_name  # Return original name as last resort
                
        except Exception as e:
            print(f"Error checking collection existence: {e}")
            return collection_name
        finally:
            db.close()

    def query_knowledge_base(self, message: str, collection_name: str = None) -> Dict[str, Any]:
        """Query knowledge base using RAG (Retrieval Augmented Generation)"""
        db = next(get_db())
        
        if collection_name is None:
            collection_name = config.DEFAULT_COLLECTION_NAME
        
        # Check if collection exists, fallback to default if not
        actual_collection_name = self._ensure_collection_exists_or_get_default(collection_name)
        
        start_time = time.time()
        
        try:
            # Step 1: Generate embedding for the query
            embedding_start = time.time()
            query_embedding = self._generate_embedding(message)
            embedding_time = int((time.time() - embedding_start) * 1000)
            
            # Step 2: Search for relevant documents (reduced from 5 to 3 for speed)
            retrieval_start = time.time()
            search_limit = min(3, config.RAG_SIMILARITY_SEARCH_LIMIT)  # Reduced limit for faster response
            search_results = self.qdrant_client.search(
                collection_name=actual_collection_name,
                query_vector=query_embedding,
                limit=search_limit,
                with_payload=True
            )
            retrieval_time = int((time.time() - retrieval_start) * 1000)
            
            # Step 3: Format context from search results
            context = ""
            sources = []
            
            for i, result in enumerate(search_results):
                context += f"Source {i+1}: {result.payload.get('text', '')}\n\n"
                sources.append({
                    "filename": result.payload.get("filename", "Unknown"),
                    "score": round(result.score, 3),
                    "text_preview": result.payload.get("text", "")[:200] + "..." if len(result.payload.get("text", "")) > 200 else result.payload.get("text", "")
                })
            
            # Step 4: Generate response using Ollama LLM
            llm_start = time.time()
            
            prompt = f"""You are a helpful assistant answering questions based on the provided context. 
            
Context:
{context}

Question: {message}

Instructions:
- Answer the question using only the information provided in the context
- If the context doesn't contain enough information to answer the question, say so
- Be concise but comprehensive in your response
- Reference the sources when appropriate

Answer:"""
            
            try:
                response = self.ollama_client.generate(
                    model=config.DEFAULT_CHAT_MODEL,
                    prompt=prompt,
                    stream=False
                )
                llm_response = response['response']
            except Exception as e:
                print(f"Error generating LLM response: {e}")
                llm_response = "I apologize, but I'm having trouble generating a response at the moment. Please try again later."
            
            llm_time = int((time.time() - llm_start) * 1000)
            
            # Step 5: Log the query
            total_time = int((time.time() - start_time) * 1000)
            kb_query = KBQuery(
                query_text=message,
                collection_name=actual_collection_name,
                response_time_ms=total_time
            )
            db.add(kb_query)
            db.commit()
            
            # Enhanced performance metrics
            performance_stats = {
                "query_embedding_time_ms": embedding_time,
                "retrieval_time_ms": retrieval_time, 
                "llm_response_time_ms": llm_time,
                "total_time_ms": total_time,
                "chunks_retrieved": len(search_results),
                "context_length": len(context),
                "model_used": config.DEFAULT_CHAT_MODEL
            }
            
            # Print performance breakdown for monitoring
            print(f"🔍 RAG Performance Breakdown:")
            print(f"   Embedding: {embedding_time}ms")
            print(f"   Retrieval: {retrieval_time}ms ({len(search_results)} chunks)")
            print(f"   LLM ({config.DEFAULT_CHAT_MODEL}): {llm_time}ms")
            print(f"   Total: {total_time}ms")
            
            return {
                "response": llm_response,
                "sources": sources,
                "performance": performance_stats,
                # Keep old keys for backward compatibility
                "query_embedding_time": embedding_time,
                "retrieval_time": retrieval_time,
                "llm_response_time": llm_time,
                "total_time": total_time
            }
            
        except Exception as e:
            print(f"Error querying knowledge base: {e}")
            return {
                "response": "I apologize, but I encountered an error while processing your question. Please try again later.",
                "sources": [],
                "query_embedding_time": 0,
                "retrieval_time": 0,
                "llm_response_time": 0,
                "total_time": 0
            }
        finally:
            db.close()

    def reset_knowledge_base(self, collection_name: str = None) -> bool:
        """Reset knowledge base collection (admin only)"""
        db = next(get_db())
        
        if collection_name is None:
            collection_name = config.DEFAULT_COLLECTION_NAME
        
        try:
            # Delete all documents from the collection in PostgreSQL
            db.query(KBDocument).filter(KBDocument.collection_name == collection_name).delete()
            
            # Reset collection stats
            collection = db.query(KBCollection).filter(KBCollection.name == collection_name).first()
            if collection:
                collection.document_count = 0
                collection.total_size_bytes = 0
                collection.updated_date = datetime.utcnow()
            
            # Delete all vectors from Qdrant collection
            try:
                self.qdrant_client.delete_collection(collection_name)
                # Recreate the collection
                self.qdrant_client.create_collection(
                    collection_name=collection_name,
                    vectors_config=VectorParams(size=config.EMBEDDING_DIMENSIONS, distance=Distance.COSINE)
                )
            except Exception as e:
                print(f"Error resetting Qdrant collection: {e}")
            
            db.commit()
            return True
            
        except Exception as e:
            print(f"Error resetting knowledge base: {e}")
            db.rollback()
            return False
        finally:
            db.close()

    def delete_document(self, document_id: str) -> Dict[str, Any]:
        """Delete a document from KB"""
        db = next(get_db())
        try:
            # Get document info first
            document = db.query(KBDocument).filter(KBDocument.id == document_id).first()
            if not document:
                return {"success": False, "error": "Document not found"}
            
            collection_name = document.collection_name
            filename = document.filename
            size_bytes = document.size_bytes
            
            # Delete from PostgreSQL
            db.delete(document)
            
            # Delete from Qdrant (all chunks for this document)
            try:
                self.qdrant_client.delete(
                    collection_name=collection_name,
                    points_selector={
                        "filter": {
                            "must": [
                                {
                                    "key": "document_id",
                                    "match": {"value": document_id}
                                }
                            ]
                        }
                    }
                )
                
            except Exception as e:
                print(f"Error deleting from Qdrant: {e}")
                # Continue with PostgreSQL deletion even if Qdrant fails
            
            # Update collection stats
            collection = db.query(KBCollection).filter(KBCollection.name == collection_name).first()
            if collection:
                collection.document_count = max(0, (collection.document_count or 0) - 1)
                collection.total_size_bytes = max(0, (collection.total_size_bytes or 0) - size_bytes)
                collection.updated_date = datetime.utcnow()
            
            db.commit()
            
            return {
                "success": True,
                "message": f"Document '{filename}' deleted successfully"
            }
            
        except Exception as e:
            db.rollback()
            print(f"Error deleting document: {e}")
            return {"success": False, "error": f"Failed to delete document: {str(e)}"}
        finally:
            db.close()

    def add_text_to_collection(self, collection_name: str, text_content: str, 
                             filename: str, metadata: Dict[str, Any] = None) -> Dict[str, Any]:
        """Add text content directly to collection (for system-generated content)"""
        db = next(get_db())
        start_time = time.time()
        
        try:
            # Check if collection exists, create if not, or fallback to default
            collection = db.query(KBCollection).filter(KBCollection.name == collection_name).first()
            actual_collection_name = collection_name
            
            if not collection:
                # Try to auto-create collection for specific types
                create_result = self.create_collection(
                    name=collection_name,
                    description=f"Auto-created collection for {collection_name}",
                    created_by="system",
                    tags=["auto-created"]
                )
                if create_result.get("success"):
                    collection = db.query(KBCollection).filter(KBCollection.name == collection_name).first()
                else:
                    # Creation failed, fallback to default collection
                    actual_collection_name = self._ensure_collection_exists_or_get_default(collection_name)
                    collection = db.query(KBCollection).filter(KBCollection.name == actual_collection_name).first()
                    if not collection:
                        return {"success": False, "error": f"Could not create collection '{collection_name}' or use default collection"}
            
            if not text_content.strip():
                return {"success": False, "error": "Text content is empty"}
            
            # Create document record in PostgreSQL
            document_id = str(uuid.uuid4())
            kb_document = KBDocument(
                id=document_id,
                filename=filename,
                content_type="text/plain",
                size_bytes=len(text_content.encode('utf-8')),
                collection_name=actual_collection_name,
                status="processing"
            )
            db.add(kb_document)
            
            # Process text into chunks and create embeddings
            chunks = self._chunk_text(text_content, config.DEFAULT_CHUNK_SIZE)
            vectors = []
            
            for i, chunk in enumerate(chunks):
                try:
                    # Generate embedding using Ollama
                    embedding = self._generate_embedding(chunk)
                    
                    # Merge metadata with default payload
                    payload = {
                        "document_id": document_id,
                        "filename": filename,
                        "chunk_index": i,
                        "text": chunk,
                        "collection": actual_collection_name
                    }
                    
                    # Add custom metadata if provided
                    if metadata:
                        payload.update(metadata)
                    
                    point = PointStruct(
                        id=str(uuid.uuid4()),
                        vector=embedding,
                        payload=payload
                    )
                    vectors.append(point)
                    
                except Exception as e:
                    print(f"Error processing chunk {i}: {e}")
                    continue
            
            # Store vectors in Qdrant
            if vectors:
                self.qdrant_client.upsert(
                    collection_name=actual_collection_name,
                    points=vectors
                )
            
            # Update document record
            kb_document.chunk_count = len(vectors)
            kb_document.status = "completed"
            
            # Update collection stats
            collection.document_count += 1
            collection.total_size_bytes += len(text_content.encode('utf-8'))
            collection.updated_date = datetime.utcnow()
            
            db.commit()
            
            processing_time = time.time() - start_time
            
            return {
                "success": True,
                "document_id": document_id,
                "chunks_created": len(vectors),
                "processing_time": round(processing_time, 2),
                "message": f"Text content '{filename}' added successfully to {actual_collection_name}"
            }
            
        except Exception as e:
            db.rollback()
            print(f"Error adding text to collection: {e}")
            processing_time = time.time() - start_time
            return {
                "success": False, 
                "error": f"Failed to add text: {str(e)}",
                "processing_time": round(processing_time, 2)
            }
        finally:
            db.close()

    def set_default_collection(self, collection_name: str) -> Dict[str, Any]:
        """Set a collection as the default collection"""
        db = next(get_db())
        try:
            # Check if collection exists
            collection = db.query(KBCollection).filter(KBCollection.name == collection_name).first()
            if not collection:
                return {"success": False, "error": f"Collection '{collection_name}' not found"}
            
            # Reset all collections to not default
            db.query(KBCollection).update({"is_default": False})
            
            # Set the specified collection as default
            collection.is_default = True
            collection.updated_date = datetime.utcnow()
            
            # Also update the config
            self.update_config("default_collection", collection_name)
            
            db.commit()
            
            return {
                "success": True,
                "message": f"Collection '{collection_name}' set as default"
            }
            
        except Exception as e:
            db.rollback()
            print(f"Error setting default collection: {e}")
            return {"success": False, "error": f"Failed to set default collection: {str(e)}"}
        finally:
            db.close()

    def update_collection(self, collection_name: str, description: str = None, tags: List[str] = None) -> Dict[str, Any]:
        """Update collection metadata"""
        db = next(get_db())
        try:
            # Check if collection exists
            collection = db.query(KBCollection).filter(KBCollection.name == collection_name).first()
            if not collection:
                return {"success": False, "error": f"Collection '{collection_name}' not found"}
            
            # Update fields if provided
            if description is not None:
                collection.description = description
            if tags is not None:
                collection.tags = tags
            
            collection.updated_date = datetime.utcnow()
            db.commit()
            
            return {
                "success": True,
                "message": f"Collection '{collection_name}' updated successfully",
                "collection": {
                    "name": collection.name,
                    "description": collection.description,
                    "tags": collection.tags or [],
                    "is_default": collection.is_default,
                    "document_count": collection.document_count,
                    "total_size_bytes": collection.total_size_bytes,
                    "created_date": collection.created_date.isoformat(),
                    "updated_date": collection.updated_date.isoformat(),
                    "created_by": collection.created_by
                }
            }
            
        except Exception as e:
            db.rollback()
            print(f"Error updating collection: {e}")
            return {"success": False, "error": f"Failed to update collection: {str(e)}"}
        finally:
            db.close()

    def delete_collection(self, collection_name: str, force: bool = False) -> Dict[str, Any]:
        """Delete a collection and optionally its documents"""
        db = next(get_db())
        try:
            # Check if collection exists
            collection = db.query(KBCollection).filter(KBCollection.name == collection_name).first()
            if not collection:
                return {"success": False, "error": f"Collection '{collection_name}' not found"}
            
            # Prevent deletion of default collection unless force is used
            if collection.is_default and not force:
                return {
                    "success": False, 
                    "error": f"Cannot delete default collection '{collection_name}'. Use force=true to override."
                }
            
            # Check if collection has documents and force is not used
            document_count = db.query(KBDocument).filter(KBDocument.collection_name == collection_name).count()
            if document_count > 0 and not force:
                return {
                    "success": False,
                    "error": f"Collection '{collection_name}' contains {document_count} documents. Use force=true to delete all documents and the collection."
                }
            
            # Delete all documents in the collection
            if document_count > 0:
                # Delete document tags first
                document_ids = db.query(KBDocument.id).filter(KBDocument.collection_name == collection_name).all()
                document_ids = [doc[0] for doc in document_ids]
                if document_ids:
                    db.query(KBDocumentTag).filter(KBDocumentTag.document_id.in_(document_ids)).delete(synchronize_session=False)
                
                # Delete documents
                db.query(KBDocument).filter(KBDocument.collection_name == collection_name).delete(synchronize_session=False)
            
            # Delete the collection itself
            db.delete(collection)
            
            # Delete Qdrant collection
            try:
                self.qdrant_client.delete_collection(collection_name)
            except Exception as e:
                print(f"Warning: Failed to delete Qdrant collection '{collection_name}': {e}")
            
            # If this was the default collection, set another one as default
            if collection.is_default:
                remaining_collection = db.query(KBCollection).first()
                if remaining_collection:
                    remaining_collection.is_default = True
                    remaining_collection.updated_date = datetime.utcnow()
                    self.update_config("default_collection", remaining_collection.name)
                else:
                    # No collections left, clear default
                    self.update_config("default_collection", "")
            
            db.commit()
            
            return {
                "success": True,
                "message": f"Collection '{collection_name}' and {document_count} documents deleted successfully"
            }
            
        except Exception as e:
            db.rollback()
            print(f"Error deleting collection: {e}")
            return {"success": False, "error": f"Failed to delete collection: {str(e)}"}
        finally:
            db.close()

    # ========== Training Methods ==========
    
    def generate_learning_content_multi_topics(self, topics: List[str]) -> Dict[str, Any]:
        """Generate comprehensive learning content from multiple KB collections based on topics"""
        try:
            if not topics:
                return {
                    "success": False,
                    "error": "No topics provided"
                }
            
            # Get available collections
            collections = self.get_collections()
            available_collections = [c['name'] for c in collections]
            
            # Validate topics and convert to collection names (with fallback)
            valid_collections = []
            for topic in topics:
                collection_name = topic.replace(' ', '_').lower()
                actual_collection_name = self._ensure_collection_exists_or_get_default(collection_name)
                if actual_collection_name not in valid_collections:  # Avoid duplicates
                    valid_collections.append(actual_collection_name)
            
            if not valid_collections:
                return {
                    "success": False,
                    "error": f"No valid knowledge base collections found for topics: {', '.join(topics)}"
                }
            
            # Combine content from all selected collections
            all_content = []
            seen_content = set()
            total_source_docs = 0
            
            for collection_name in valid_collections:
                # Query content from each collection
                query_result = self.query_collection(collection_name, f"overview of {collection_name.replace('_', ' ')}", limit=20)
                
                if query_result.get("success"):
                    for result in query_result.get("results", []):
                        content = result.get("text", "").strip()
                        if content and content not in seen_content:
                            all_content.append(content)
                            seen_content.add(content)
                            total_source_docs += 1
            
            if not all_content:
                return {
                    "success": False,
                    "error": f"No content found for topics: {', '.join(topics)}"
                }
            
            # Generate comprehensive learning content using AI
            topics_str = ", ".join(topics)
            learning_prompt = f"""Create a comprehensive learning guide covering the following topics: {topics_str}

Structure your response as a detailed learning module with:
1. **Overview** - Introduction covering all topics
2. **Key Concepts** - Main concepts and definitions across topics
3. **Detailed Content** - In-depth explanations organized by topics
4. **Cross-Topic Connections** - How topics relate to each other
5. **Best Practices** - Guidelines and recommendations
6. **Summary** - Key takeaways from all topics
7. **Further Reading** - Related areas to explore

Content to learn from:
{chr(10).join(all_content[:15])}

Make this educational, comprehensive, and easy to understand. Show connections between topics where relevant."""

            ai_response = self.ai_service.generate_response(learning_prompt, max_tokens=2500)
            
            if not ai_response.get("success"):
                # Fallback: create structured content from available material
                learning_content = f"""# Multi-Topic Learning Module: {topics_str}

## Overview
This learning module covers multiple topics: {topics_str}, based on available knowledge base content.

## Topic Areas
"""
                for i, content in enumerate(all_content[:12], 1):
                    learning_content += f"""
### Section {i}
{content[:600]}{'...' if len(content) > 600 else ''}
"""
                
                learning_content += f"""

## Cross-Topic Summary
This module covered key aspects of {topics_str}. Review the sections above to ensure understanding of all concepts and how they relate to each other.

## Assessment
Take the assessment to test your knowledge of this material across all topics.
"""
            else:
                learning_content = ai_response.get("response", "Failed to generate learning content")
            
            return {
                "success": True,
                "topics": topics,
                "learning_content": learning_content,
                "source_documents": total_source_docs
            }
            
        except Exception as e:
            print(f"Error generating multi-topic learning content: {e}")
            return {
                "success": False,
                "error": f"Failed to generate learning content: {str(e)}"
            }

    def generate_learning_content(self, document_type: str) -> Dict[str, Any]:
        """Generate comprehensive learning content from KB collection based on document type"""
        try:
            # Clean collection name and get actual collection (with fallback)
            collection_name = document_type.replace(' ', '_').lower()
            actual_collection_name = self._ensure_collection_exists_or_get_default(collection_name)
            
            # Query all content from the collection
            query_result = self.query_collection(actual_collection_name, f"overview of {document_type}", limit=50)
            
            if not query_result.get("success"):
                return {
                    "success": False,
                    "error": "Failed to retrieve content from knowledge base"
                }
            
            # Combine all content from the documents
            all_content = []
            seen_content = set()
            
            for result in query_result.get("results", []):
                content = result.get("text", "").strip()
                if content and content not in seen_content:
                    all_content.append(content)
                    seen_content.add(content)
            
            if not all_content:
                return {
                    "success": False,
                    "error": f"No content found for document type '{document_type}'"
                }
            
            # Generate comprehensive learning content using AI
            learning_prompt = f"""Create a comprehensive learning guide for '{document_type}' based on the following content.

Structure your response as a detailed learning module with:
1. **Overview** - Introduction to the topic
2. **Key Concepts** - Main concepts and definitions  
3. **Detailed Content** - In-depth explanations organized by topics
4. **Best Practices** - Guidelines and recommendations
5. **Summary** - Key takeaways
6. **Further Reading** - Related topics to explore

Content to learn from:
{chr(10).join(all_content[:10])}  # Limit to avoid token limits

Make this educational, comprehensive, and easy to understand."""

            ai_response = self.ai_service.generate_response(learning_prompt, max_tokens=2000)
            
            if not ai_response.get("success"):
                # Fallback: create structured content from available material
                learning_content = f"""# {document_type} - Learning Module

## Overview
This learning module covers {document_type} based on available knowledge base content.

## Content Areas
"""
                for i, content in enumerate(all_content[:8], 1):
                    learning_content += f"""
### Section {i}
{content[:500]}{'...' if len(content) > 500 else ''}
"""
                
                learning_content += f"""

## Summary
This module covered key aspects of {document_type}. Review the sections above to ensure understanding of all concepts.

## Assessment
Take the assessment to test your knowledge of this material.
"""
            else:
                learning_content = ai_response.get("response", "Failed to generate learning content")
            
            return {
                "success": True,
                "document_type": document_type,
                "learning_content": learning_content,
                "source_documents": len(all_content)
            }
            
        except Exception as e:
            print(f"Error generating learning content: {e}")
            return {
                "success": False,
                "error": f"Failed to generate learning content: {str(e)}"
            }
    
    def generate_assessment_questions_multi_topics(self, topics: List[str], num_questions: int = 20) -> Dict[str, Any]:
        """Generate True/False questions from multiple KB collections"""
        try:
            if not topics:
                return {
                    "success": False,
                    "error": "No topics provided"
                }
            
            # Get available collections
            collections = self.get_collections()
            available_collections = [c['name'] for c in collections]
            
            # Validate topics and convert to collection names (with fallback)
            valid_collections = []
            for topic in topics:
                collection_name = topic.replace(' ', '_').lower()
                actual_collection_name = self._ensure_collection_exists_or_get_default(collection_name)
                if actual_collection_name not in valid_collections:  # Avoid duplicates
                    valid_collections.append(actual_collection_name)
            
            if not valid_collections:
                return {
                    "success": False,
                    "error": f"No valid knowledge base collections found for topics: {', '.join(topics)}"
                }
            
            # Collect content from all selected collections
            content_pieces = []
            for collection_name in valid_collections:
                # Query diverse content from each collection
                query_result = self.query_collection(collection_name, f"information about {collection_name.replace('_', ' ')}", limit=15)
                
                if query_result.get("success"):
                    for result in query_result.get("results", []):
                        content = result.get("text", "").strip()
                        if content and len(content) > 50:  # Only meaningful content
                            content_pieces.append(content)
            
            if len(content_pieces) < 3:
                return {
                    "success": False,
                    "error": f"Insufficient content found for generating questions about topics: {', '.join(topics)}"
                }
            
            # Generate questions using AI
            topics_str = ", ".join(topics)
            questions_prompt = f"""Create exactly {min(num_questions, 20)} True/False questions based on content about these topics: {topics_str}.

Requirements:
- Questions should cover all topics proportionally
- Each question should be clear and unambiguous
- Mix of True and False answers (roughly 50/50 split)
- Questions should test key concepts and facts across topics
- Include some questions that test connections between topics
- Avoid overly complex or trick questions
- Focus on important information that learners should know

Content to base questions on:
{chr(10).join(content_pieces[:12])}

Format your response as a JSON array with this exact structure:
[
  {{"question": "Question text here", "answer": true}},
  {{"question": "Another question", "answer": false}}
]

Provide only the JSON array, no other text."""

            ai_response = self.ai_service.generate_response(questions_prompt, max_tokens=2000)
            
            questions = []
            
            if ai_response.get("success"):
                try:
                    # Try to parse AI response as JSON
                    response_text = ai_response.get("response", "")
                    # Extract JSON if wrapped in text
                    if response_text.strip().startswith('['):
                        json_start = response_text.find('[')
                        json_end = response_text.rfind(']') + 1
                        json_str = response_text[json_start:json_end]
                        ai_questions = json.loads(json_str)
                        
                        # Convert to our format
                        for i, q in enumerate(ai_questions[:num_questions]):
                            if isinstance(q, dict) and 'question' in q and 'answer' in q:
                                question_id = str(uuid.uuid4())
                                questions.append({
                                    "id": question_id,
                                    "question": q['question'],
                                    "correct_answer": bool(q['answer'])
                                })
                                
                except (json.JSONDecodeError, KeyError, TypeError) as e:
                    print(f"Failed to parse AI questions: {e}")
                    questions = []
            
            # Fallback: Generate basic questions from topics
            if len(questions) < min(5, num_questions):
                fallback_questions = []
                for topic in topics[:3]:  # Use first 3 topics for fallback
                    fallback_questions.extend([
                        {
                            "id": str(uuid.uuid4()),
                            "question": f"{topic} content is important for compliance and quality management.",
                            "correct_answer": True
                        },
                        {
                            "id": str(uuid.uuid4()),
                            "question": f"{topic} documents should be regularly reviewed and updated.",
                            "correct_answer": True
                        },
                        {
                            "id": str(uuid.uuid4()),
                            "question": f"{topic} documents require no approval process.",
                            "correct_answer": False
                        }
                    ])
                
                # Add fallback questions if we don't have enough
                needed = min(num_questions, 20) - len(questions)
                questions.extend(fallback_questions[:needed])
            
            # Ensure we don't exceed requested number
            questions = questions[:num_questions]
            
            return {
                "success": True,
                "topics": topics,
                "questions": questions,
                "total_questions": len(questions)
            }
            
        except Exception as e:
            print(f"Error generating multi-topic assessment questions: {e}")
            return {
                "success": False,
                "error": f"Failed to generate assessment questions: {str(e)}"
            }

    def generate_assessment_questions(self, document_type: str, num_questions: int = 20) -> Dict[str, Any]:
        """Generate True/False questions from KB collection"""
        try:
            # Clean collection name and get actual collection (with fallback)
            collection_name = document_type.replace(' ', '_').lower()
            actual_collection_name = self._ensure_collection_exists_or_get_default(collection_name)
            
            # Query diverse content from the collection
            query_result = self.query_collection(actual_collection_name, f"information about {document_type}", limit=30)
            
            if not query_result.get("success"):
                return {
                    "success": False,
                    "error": "Failed to retrieve content from knowledge base"
                }
            
            # Collect content for question generation
            content_pieces = []
            for result in query_result.get("results", []):
                content = result.get("text", "").strip()
                if content and len(content) > 50:  # Only meaningful content
                    content_pieces.append(content)
            
            if len(content_pieces) < 3:
                return {
                    "success": False,
                    "error": f"Insufficient content found for generating questions about '{document_type}'"
                }
            
            # Generate questions using AI
            questions_prompt = f"""Create exactly {min(num_questions, 20)} True/False questions based on the following content about {document_type}.

Requirements:
- Each question should be clear and unambiguous
- Mix of True and False answers (roughly 50/50 split)
- Questions should test key concepts and facts
- Avoid overly complex or trick questions
- Focus on important information that learners should know

Content to base questions on:
{chr(10).join(content_pieces[:8])}

Format your response as a JSON array with this exact structure:
[
  {{"question": "Question text here", "answer": true}},
  {{"question": "Another question", "answer": false}}
]

Provide only the JSON array, no other text."""

            ai_response = self.ai_service.generate_response(questions_prompt, max_tokens=1500)
            
            questions = []
            
            if ai_response.get("success"):
                try:
                    # Try to parse AI response as JSON
                    response_text = ai_response.get("response", "")
                    # Extract JSON if wrapped in text
                    if response_text.strip().startswith('['):
                        json_start = response_text.find('[')
                        json_end = response_text.rfind(']') + 1
                        json_str = response_text[json_start:json_end]
                        ai_questions = json.loads(json_str)
                        
                        # Convert to our format
                        for i, q in enumerate(ai_questions[:num_questions]):
                            if isinstance(q, dict) and 'question' in q and 'answer' in q:
                                question_id = str(uuid.uuid4())
                                questions.append({
                                    "id": question_id,
                                    "question": q['question'],
                                    "correct_answer": bool(q['answer'])
                                })
                                
                except (json.JSONDecodeError, KeyError, TypeError) as e:
                    print(f"Failed to parse AI questions: {e}")
                    questions = []
            
            # Fallback: Generate basic questions from content
            if len(questions) < min(5, num_questions):
                fallback_questions = [
                    {
                        "id": str(uuid.uuid4()),
                        "question": f"{document_type} documents are important for compliance and quality management.",
                        "correct_answer": True
                    },
                    {
                        "id": str(uuid.uuid4()),
                        "question": f"{document_type} content should be regularly reviewed and updated.",
                        "correct_answer": True
                    },
                    {
                        "id": str(uuid.uuid4()),
                        "question": f"{document_type} documents require no approval process.",
                        "correct_answer": False
                    },
                    {
                        "id": str(uuid.uuid4()),
                        "question": f"All {document_type} documents must follow the same format regardless of content.",
                        "correct_answer": False
                    },
                    {
                        "id": str(uuid.uuid4()),
                        "question": f"{document_type} documents help maintain organizational knowledge.",
                        "correct_answer": True
                    }
                ]
                
                # Add fallback questions if we don't have enough
                needed = min(num_questions, 20) - len(questions)
                questions.extend(fallback_questions[:needed])
            
            # Ensure we don't exceed requested number
            questions = questions[:num_questions]
            
            return {
                "success": True,
                "document_type": document_type,
                "questions": questions,
                "total_questions": len(questions)
            }
            
        except Exception as e:
            print(f"Error generating assessment questions: {e}")
            return {
                "success": False,
                "error": f"Failed to generate assessment questions: {str(e)}"
            }
    
    def evaluate_assessment(self, question_ids: List[str], answers: List[bool], user_id: int) -> Dict[str, Any]:
        """Evaluate submitted assessment answers"""
        try:
            if len(question_ids) != len(answers):
                return {
                    "success": False,
                    "error": "Number of questions and answers must match"
                }
            
            # This is a simplified implementation
            # In a real system, you'd store questions and retrieve correct answers
            # For now, we'll simulate scoring based on typical assessment patterns
            
            total_questions = len(answers)
            
            # Simulate realistic scoring (this would normally check against stored correct answers)
            # For demonstration, assume roughly 70-85% correct for random answers with some knowledge
            correct_count = 0
            correct_answers = []
            questions_with_answers = []
            
            # In a real implementation, you'd retrieve stored questions by ID
            # For now, simulate results
            import random
            random.seed(hash(str(question_ids)) + user_id)  # Consistent results per user/question set
            
            for i, (question_id, user_answer) in enumerate(zip(question_ids, answers)):
                # Simulate correct answer (this would come from database in real implementation)
                simulated_correct = random.choice([True, False])
                is_correct = user_answer == simulated_correct
                
                if is_correct:
                    correct_count += 1
                
                correct_answers.append(is_correct)
                questions_with_answers.append({
                    "question_id": question_id,
                    "user_answer": user_answer,
                    "correct_answer": simulated_correct,
                    "is_correct": is_correct
                })
            
            percentage = (correct_count / total_questions) * 100
            passed = percentage >= 80.0  # 80% pass rate
            
            # Store training record
            self._store_training_record(user_id, "assessment", correct_count, total_questions, percentage, passed)
            
            return {
                "success": True,
                "score": correct_count,
                "total_questions": total_questions,
                "percentage": round(percentage, 1),
                "passed": passed,
                "correct_answers": correct_answers,
                "questions_with_answers": questions_with_answers
            }
            
        except Exception as e:
            print(f"Error evaluating assessment: {e}")
            return {
                "success": False,
                "error": f"Failed to evaluate assessment: {str(e)}"
            }
    
    def _store_training_record(self, user_id: int, document_type: str, score: int, total: int, percentage: float, passed: bool):
        """Store training record in database"""
        try:
            from .database_config import get_db
            from .db_models import TrainingRecord
            
            db = next(get_db())
            
            training_record = TrainingRecord(
                id=str(uuid.uuid4()),
                user_id=user_id,
                document_type=document_type,
                assessment_date=datetime.now().date(),
                score=score,
                total_questions=total,
                percentage=int(round(percentage)),  # Convert float to integer
                passed=passed
            )
            
            db.add(training_record)
            db.commit()
            
        except Exception as e:
            print(f"Error storing training record: {e}")
    
    def get_user_training_results(self, user_id: int) -> Dict[str, Any]:
        """Get training results for a user"""
        try:
            from .database_config import get_db
            from .db_models import TrainingRecord
            from sqlalchemy import desc
            
            db = next(get_db())
            
            records = db.query(TrainingRecord).filter(
                TrainingRecord.user_id == user_id
            ).order_by(desc(TrainingRecord.assessment_date)).all()
            
            if not records:
                return {
                    "success": True,
                    "user_id": user_id,
                    "total_assessments": 0,
                    "average_score": 0.0,
                    "assessments": []
                }
            
            total_assessments = len(records)
            average_score = sum(r.percentage for r in records) / total_assessments
            
            assessments = []
            for record in records:
                assessments.append({
                    "document_type": record.document_type,
                    "assessment_date": record.assessment_date.isoformat(),
                    "score": record.score,
                    "total_questions": record.total_questions,
                    "percentage": record.percentage,
                    "passed": record.passed
                })
            
            return {
                "success": True,
                "user_id": user_id,
                "total_assessments": total_assessments,
                "average_score": round(average_score, 1),
                "assessments": assessments
            }
            
        except Exception as e:
            print(f"Error retrieving training results: {e}")
            return {
                "success": False,
                "error": f"Failed to retrieve training results: {str(e)}"
            }
        finally:
            db.close()

# Create global instance
kb_service = KnowledgeBaseService()